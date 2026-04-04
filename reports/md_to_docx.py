"""
md_to_docx.py
=============
Convert reports/midterm_report.md → reports/midterm_report.docx

Handles:
  - H1 / H2 / H3 / H4 headings  (mapped to Title / Heading 1-3)
  - Inline bold (**), italic (*), inline-code (`) — including mixed
  - Fenced code blocks  (``` ... ```)
  - Markdown tables
  - Image embeds  (![alt](charts/foo.png))
  - Bullet lists (- item)
  - Numbered lists (1. item)
  - [INSERT …] screenshot placeholders  → highlighted grey box
  - Horizontal rules (---) → thin separator line
  - Automatic Table of Contents field (update with Ctrl+A → F9 in Word)
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPORT_DIR = Path(__file__).parent
CHARTS_DIR  = REPORT_DIR / "charts"
MD_FILE     = REPORT_DIR / "midterm_report.md"
DOCX_FILE   = REPORT_DIR / "midterm_report.docx"

# Map markdown image paths → actual file paths
CHART_MAP: dict[str, Path] = {
    "charts/feature_missing_rates.png":    CHARTS_DIR / "feature_missing_rates.png",
    "charts/genre_count_distribution.png": CHARTS_DIR / "genre_count_distribution.png",
    "charts/rule_auc_comparison.png":      CHARTS_DIR / "rule_auc_comparison.png",
    "charts/fallback_usage.png":           CHARTS_DIR / "fallback_usage.png",
}

# Colour constants
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)   # Word default heading blue
CODE_GREY  = "F2F2F2"                       # code block background
RULE_GREY  = "AAAAAA"                       # horizontal rule colour
TOC_FILL   = "EBF3FB"                       # TOC box background


# ===========================================================================
# Low-level XML helpers
# ===========================================================================

def _set_shading(paragraph, fill_hex: str) -> None:
    """Apply a solid background colour to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _set_run_shading(run, fill_hex: str) -> None:
    """Apply a background shading to an individual run (inline code)."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    rPr.append(shd)


def _add_bottom_border(paragraph, colour: str = RULE_GREY, sz: str = "6") -> None:
    """Add a bottom border to a paragraph (used for --- rules)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_toc_field(document: Document) -> None:
    """Insert a TOC field that Word will populate on Ctrl+A → F9."""
    heading = document.add_paragraph("Table of Contents", style="Heading 1")
    heading.runs[0].font.color.rgb = DARK_BLUE

    p = document.add_paragraph()
    run = p.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "(Right-click → Update Field to populate)"
    placeholder_run = OxmlElement("w:r")
    placeholder_run.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder_run)
    run._r.append(end)

    # Light blue background for the TOC block so it's visible before update
    _set_shading(p, TOC_FILL)

    # Page break after TOC
    document.add_page_break()


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Shade a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


# ===========================================================================
# Inline formatting parser
# ===========================================================================
# Tokenises text into segments: (text, bold, italic, is_code)
# Handles **bold**, *italic*, `code`, and combinations.
_INLINE_RE = re.compile(
    r"(\*\*[^*]+?\*\*"   # **bold**
    r"|\*[^*]+?\*"       # *italic*
    r"|`[^`]+?`)"        # `code`
)


def _parse_inline(text: str) -> list[tuple[str, bool, bool, bool]]:
    """Return [(text, bold, italic, code), ...] tokens."""
    segments: list[tuple[str, bool, bool, bool]] = []
    last = 0
    for m in _INLINE_RE.finditer(text):
        # Text before match
        if m.start() > last:
            segments.append((text[last:m.start()], False, False, False))
        tok = m.group(0)
        if tok.startswith("**"):
            segments.append((tok[2:-2], True, False, False))
        elif tok.startswith("*"):
            segments.append((tok[1:-1], False, True, False))
        else:  # `code`
            segments.append((tok[1:-1], False, False, True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False, False, False))
    return segments


def _add_inline_runs(
    paragraph,
    text: str,
    base_bold: bool = False,
    base_italic: bool = False,
    base_code: bool = False,
    font_size: int | None = None,
) -> None:
    """Parse *text* and add appropriately formatted runs to *paragraph*."""
    for raw, bold, italic, code in _parse_inline(text):
        if not raw:
            continue
        run = paragraph.add_run(raw)
        run.bold   = base_bold  or bold
        run.italic = base_italic or italic
        if base_code or code:
            run.font.name = "Courier New"
            run.font.size = Pt(font_size or 9)
            _set_run_shading(run, CODE_GREY)
        elif font_size:
            run.font.size = Pt(font_size)


# ===========================================================================
# Document-level helpers
# ===========================================================================

def _apply_heading_style(paragraph, colour: RGBColor = DARK_BLUE) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = colour


def _add_code_block(document: Document, lines: list[str]) -> None:
    """Render a fenced code block with monospace font and grey background."""
    # One paragraph per line so each line gets the grey background.
    for line in lines:
        p = document.add_paragraph()
        _set_shading(p, CODE_GREY)
        # Preserve indentation with a left indent
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    # Small gap after block
    gap = document.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after  = Pt(2)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    """Render a markdown table into a Word table."""
    if not rows:
        return

    # Find separator row (row of ---) and strip it
    data_rows = [r for r in rows if not all(re.match(r"^[-: ]+$", c) for c in r)]
    if not data_rows:
        return

    ncols = max(len(r) for r in data_rows)
    # Pad short rows
    data_rows = [r + [""] * (ncols - len(r)) for r in data_rows]

    tbl = document.add_table(rows=len(data_rows), cols=ncols)
    tbl.style = "Table Grid"

    for ri, row in enumerate(data_rows):
        tr = tbl.rows[ri]
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            is_header = ri == 0
            if is_header:
                _set_cell_shading(cell, "1F497D")  # dark blue header
                run = p.add_run(cell_text.strip())
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            else:
                _set_cell_shading(cell, "FFFFFF" if ri % 2 == 1 else "EBF3FB")
                _add_inline_runs(p, cell_text.strip(), font_size=10)

    document.add_paragraph()  # spacing after table


def _add_image(document: Document, img_path: str, alt: str) -> None:
    """Embed an image centred, with an italic caption below."""
    resolved = CHART_MAP.get(img_path)
    if resolved is None or not resolved.exists():
        p = document.add_paragraph(f"[IMAGE NOT FOUND: {img_path}]")
        p.runs[0].italic = True
        return

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(resolved), width=Inches(5.8))

    caption = document.add_paragraph(alt)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    caption.paragraph_format.space_after = Pt(8)


def _add_placeholder(document: Document, text: str) -> None:
    """Render [INSERT ...] as a shaded, centred, italic notice."""
    p = document.add_paragraph()
    _set_shading(p, "FFF3CD")   # amber warning background
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
    run.font.size = Pt(10)


def _add_hr(document: Document) -> None:
    """Thin horizontal separator line."""
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    _add_bottom_border(p)


# ===========================================================================
# Main parser / converter
# ===========================================================================

def _parse_table_row(line: str) -> list[str]:
    """Split a markdown table row into cell strings."""
    line = line.strip().lstrip("|").rstrip("|")
    return [c.strip() for c in line.split("|")]


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # ── Document-wide defaults ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Narrow margins for more content width
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # ── Read markdown ────────────────────────────────────────────────────────
    lines = md_path.read_text(encoding="utf-8").splitlines()

    # ── State ────────────────────────────────────────────────────────────────
    in_code_block   = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_table        = False
    toc_inserted    = False

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if table_rows:
            _add_table(doc, table_rows)
        in_table   = False
        table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Fenced code block toggle ─────────────────────────────────────
        if stripped.startswith("```"):
            if in_code_block:
                _add_code_block(doc, code_lines)
                code_lines   = []
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)   # preserve indentation
            i += 1
            continue

        # ── Markdown table rows ──────────────────────────────────────────
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(_parse_table_row(stripped))
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── Empty line ───────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────────
        if stripped == "---":
            _add_hr(doc)
            i += 1
            continue

        # ── Image ────────────────────────────────────────────────────────
        img_m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_m:
            alt, path = img_m.group(1), img_m.group(2)
            _add_image(doc, path, alt)
            i += 1
            continue

        # ── [INSERT …] placeholder ───────────────────────────────────────
        if stripped.startswith("[INSERT"):
            _add_placeholder(doc, stripped)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────
        heading_m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_m:
            level = len(heading_m.group(1))
            text  = heading_m.group(2).strip()

            if level == 1:
                # Document title — use Title style
                p = doc.add_paragraph()
                p.style = doc.styles["Title"]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = DARK_BLUE

                # Insert TOC immediately after the title
                if not toc_inserted:
                    doc.add_paragraph()   # small gap
                    _add_toc_field(doc)
                    toc_inserted = True

            else:
                # H2 → Heading 1, H3 → Heading 2, H4 → Heading 3
                word_level = level - 1
                style_name = f"Heading {min(word_level, 9)}"
                p = doc.add_paragraph(style=style_name)
                p.clear()   # remove auto-added run so we control formatting

                # Strip any trailing " ← BEST" annotation — keep it as a run
                annotation = ""
                if "←" in text:
                    parts = text.split("←", 1)
                    text, annotation = parts[0].rstrip(), "← " + parts[1].strip()

                _add_inline_runs(p, text, base_bold=(word_level == 1))
                if annotation:
                    run = p.add_run("  " + annotation)
                    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                    run.bold = False
                    run.font.size = Pt(10)

                _apply_heading_style(p)

            i += 1
            continue

        # ── Bullet list item ─────────────────────────────────────────────
        bullet_m = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, bullet_m.group(1))
            i += 1
            continue

        # ── Numbered list item ───────────────────────────────────────────
        num_m = re.match(r"^\d+\.\s+(.*)", stripped)
        if num_m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, num_m.group(1))
            i += 1
            continue

        # ── Normal paragraph (may span multiple physical lines) ──────────
        # Gather consecutive non-blank, non-special lines into one paragraph
        para_lines = []
        while i < len(lines):
            ln = lines[i].strip()
            if (not ln
                or ln == "---"
                or ln.startswith("#")
                or ln.startswith("|")
                or ln.startswith("```")
                or re.match(r"!\[", ln)
                or ln.startswith("[INSERT")
                or re.match(r"^[-*]\s", ln)
                or re.match(r"^\d+\.\s", ln)):
                break
            para_lines.append(ln)
            i += 1

        text = " ".join(para_lines)
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _add_inline_runs(p, text)

    # Flush any trailing table
    if in_table:
        flush_table()

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert(MD_FILE, DOCX_FILE)
